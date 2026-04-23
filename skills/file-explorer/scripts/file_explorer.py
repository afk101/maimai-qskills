#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件内容探索引擎

功能说明：
    1. 搜索模式：读取 PDF/DOCX/Excel 文件的文本内容，
       通过正则表达式匹配搜索包含指定关键词的文件，
       输出匹配结果（JSON 格式）。
    2. 读取模式：根据文件路径数组，返回每个文件的完整文本内容（JSON 格式）。

支持的文件格式：
    - PDF (.pdf)
    - Word (.docx)
    - Excel (.xlsx, .xls)

用法：
    # 搜索关键词
    python3 file_explorer.py search --dirs "/path1,/path2" --pattern "正则表达式"

    # 读取单个文件内容
    python3 file_explorer.py read --files "/path/to/file.pdf"

    # 读取多个文件内容（逗号分隔）
    python3 file_explorer.py read --files "/path/to/a.pdf,/path/to/b.docx"

    # 读取多个文件内容（JSON 数组）
    python3 file_explorer.py read --files-json '["/path/to/a.pdf", "/path/to/b.docx"]'

依赖：
    pip3 install pdfplumber python-docx openpyxl
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

# 导入索引引擎
from file_index_sqlite import FileIndexSQLite as FileIndex
from file_lock import file_lock
from background_indexer import BackgroundIndexer


def check_dependencies():
    """
    检查所需的 Python 依赖是否已安装。

    @returns {dict} 依赖检查结果，key 为包名，value 为是否安装
    """
    dependencies = {
        "pdfplumber": False,
        "docx": False,
        "openpyxl": False,
    }
    try:
        import pdfplumber  # noqa: F401
        dependencies["pdfplumber"] = True
    except ImportError:
        pass
    try:
        import docx  # noqa: F401
        dependencies["docx"] = True
    except ImportError:
        pass
    try:
        import openpyxl  # noqa: F401
        dependencies["openpyxl"] = True
    except ImportError:
        pass
    return dependencies


def parse_pdf(filepath):
    """
    解析 PDF 文件，提取所有页面的文本内容。

    @param {str} filepath - PDF 文件的绝对路径
    @returns {str} 提取的文本内容，多页以换行连接
    @raises Exception 当文件无法读取时抛出异常
    """
    import pdfplumber

    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(filepath):
    """
    解析 DOCX 文件，提取所有段落和表格的文本内容。

    @param {str} filepath - DOCX 文件的绝对路径
    @returns {str} 提取的文本内容，多段落以换行连接
    @raises Exception 当文件无法读取时抛出异常
    """
    import docx

    doc = docx.Document(filepath)
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # 同时提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)
    return "\n".join(text_parts)


def parse_excel(filepath):
    """
    解析 Excel 文件，提取所有 sheet 中所有单元格的文本内容。
    返回结构化内容：每个 sheet 作为一个块，每行以 Tab 分隔。

    @param {str} filepath - Excel 文件的绝对路径
    @returns {str} 提取的文本内容
    @raises Exception 当文件无法读取时抛出异常
    """
    import openpyxl

    workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheet_parts = []
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        row_parts = []
        for row in sheet.iter_rows(values_only=True):
            # 只保留有内容的行
            cells = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in cells):
                row_parts.append("\t".join(cells))
        if row_parts:
            sheet_parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(row_parts))
    workbook.close()
    return "\n\n".join(sheet_parts)


# 文件扩展名与解析函数的映射关系
PARSER_MAP = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
}


def read_file_content(filepath):
    """
    读取单个文件的完整文本内容。

    @param {str} filepath - 文件的绝对路径
    @returns {dict} 包含 filepath/content/size_chars/error 的字典
                    成功时 error 为 None，失败时 content 为 None
    """
    filepath = os.path.expanduser(filepath.strip())

    if not os.path.exists(filepath):
        return {
            "filepath": filepath,
            "content": None,
            "size_chars": 0,
            "error": f"文件不存在: {filepath}",
        }

    ext = os.path.splitext(filepath)[1].lower()
    parser = PARSER_MAP.get(ext)

    if parser is None:
        return {
            "filepath": filepath,
            "content": None,
            "size_chars": 0,
            "error": f"不支持的文件格式: {ext}（支持: {', '.join(PARSER_MAP.keys())}）",
        }

    try:
        content = parser(filepath)
        return {
            "filepath": filepath,
            "content": content if content else "",
            "size_chars": len(content) if content else 0,
            "error": None,
        }
    except Exception as e:
        return {
            "filepath": filepath,
            "content": None,
            "size_chars": 0,
            "error": f"解析失败: {str(e)}",
        }


def read_multiple_files(filepaths):
    """
    批量读取多个文件的完整文本内容。

    @param {list} filepaths - 文件路径列表
    @returns {dict} 包含 results/total/succeeded/failed 的字典
    """
    results = []
    total = len(filepaths)

    for index, filepath in enumerate(filepaths):
        print(
            f"\r[进度] {index + 1}/{total} 正在读取: {os.path.basename(filepath)}...",
            end="",
            file=sys.stderr,
        )
        result = read_file_content(filepath)
        results.append(result)

    print("", file=sys.stderr)

    succeeded = sum(1 for r in results if r["error"] is None)
    failed = total - succeeded

    return {
        "results": results,
        "total": total,
        "succeeded": succeeded,
        "failed": failed,
    }


def search_in_file(filepath, pattern):
    """
    在单个文件中执行正则匹配搜索。

    @param {str} filepath - 文件的绝对路径
    @param {str} pattern - 正则表达式模式字符串
    @returns {dict|None} 匹配结果字典，包含 filepath/matches/match_count；无匹配返回 None
    """
    ext = os.path.splitext(filepath)[1].lower()
    parser = PARSER_MAP.get(ext)
    if parser is None:
        return None

    try:
        text = parser(filepath)
        if not text:
            return None

        matches = re.findall(pattern, text)
        if matches:
            return {
                "filepath": filepath,
                "matches": matches[:10],  # 最多返回前10个匹配项，避免输出过多
                "match_count": len(matches),
            }
    except Exception as e:
        # 输出解析错误到标准错误流，不影响正常结果输出
        print(f"[警告] 解析文件失败: {filepath} - {str(e)}", file=sys.stderr)
    return None


def collect_files(directories, extensions):
    """
    遍历指定目录，收集所有符合扩展名的文件路径。

    @param {list} directories - 要遍历的目录路径列表
    @param {list} extensions - 要匹配的文件扩展名列表（如 ['.pdf', '.docx']）
    @returns {list} 所有符合条件的文件绝对路径列表
    """
    files = []
    for directory in directories:
        if not os.path.isdir(directory):
            print(f"[警告] 目录不存在，已跳过: {directory}", file=sys.stderr)
            continue
        for root, _dirs, filenames in os.walk(directory):
            for filename in filenames:
                ext = os.path.splitext(filename)[1].lower()
                if ext in extensions:
                    files.append(os.path.join(root, filename))
    return files


def search_in_directories(directories, pattern, extensions, index_path=None):
    """
    在多个目录中搜索匹配正则表达式的文件。

    两阶段搜索：
        1. 文件名匹配（快速路径，good case: 简历场景）
        2. 内容匹配（使用索引，避免重复解析）

    流式处理：找到结果立即返回，不等待全部索引完成。

    @param {list} directories - 要遍历的目录路径列表
    @param {str} pattern - 正则表达式模式字符串
    @param {list} extensions - 要匹配的文件扩展名列表
    @param {str} index_path - 索引文件路径，None 表示不使用索引
    @returns {dict} 搜索结果，包含 results/total_files/matched_files/errors 字段
    """
    errors = []
    files = collect_files(directories, extensions)
    total_files = len(files)

    # 阶段 1: 文件名匹配（极快）
    print("[阶段1] 文件名匹配...", file=sys.stderr)
    filename_results = []
    try:
        regex = re.compile(pattern)
        for filepath in files:
            filename = os.path.basename(filepath)
            name_without_ext = os.path.splitext(filename)[0]
            if regex.search(name_without_ext):
                filename_results.append({
                    "filepath": filepath,
                    "filename": filename,
                    "match_type": "filename",
                    "match_count": 1,
                    "matches": [name_without_ext]
                })
    except re.error as e:
        errors.append(f"无效的正则表达式: {str(e)}")
        return {
            "results": [],
            "total_files": total_files,
            "matched_files": 0,
            "errors": errors,
        }

    if filename_results:
        print(f"[阶段1] 文件名匹配到 {len(filename_results)} 个文件，立即返回", file=sys.stderr)

        # 启动后台索引（所有文件都未索引）
        if index_path:
            lock_file = os.path.join(os.path.dirname(index_path), ".lock")
            pending_path = os.path.join(os.path.dirname(index_path), "pending_index.json")
            file_index_bg = FileIndex(index_path, lock_file)
            background_indexer = BackgroundIndexer(
                file_index_bg,
                lock_file,
                pending_path,
                read_file_content_func=read_file_content
            )
            background_indexer.start_background_indexing(
                files,
                start_from_index=0,
                search_id=f"filename-{pattern}"
            )

        # 文件名匹配到结果，直接返回，不继续内容匹配
        return {
            "results": filename_results,
            "total_files": total_files,
            "matched_files": len(filename_results),
            "filename_matches": len(filename_results),
            "content_matches": 0,
            "errors": errors,
        }

    # 阶段 2: 内容匹配（使用索引）
    print("[阶段2] 内容匹配...", file=sys.stderr)
    content_results = []

    # 使用索引
    if index_path:
        # 初始化锁文件路径
        lock_file = os.path.join(os.path.dirname(index_path), ".lock")

        # 初始化 FileIndex（带锁文件）
        file_index = FileIndex(index_path, lock_file)

        # 检查待索引队列（用于恢复）
        pending_path = os.path.join(os.path.dirname(index_path), "pending_index.json")
        background_indexer = BackgroundIndexer(
            file_index,
            lock_file,
            pending_path,
            read_file_content_func=read_file_content
        )

        pending = background_indexer.load_pending_queue()
        if pending:
            print(
                f"[信息] 恢复待索引队列 ({len(pending.get('pending_files', []))} 个文件)",
                file=sys.stderr
            )

        for index, filepath in enumerate(files):
            # 进度
            progress = int((index + 1) / total_files * 100) if total_files > 0 else 0
            print(
                f"\r[阶段2] {index + 1}/{total_files} ({progress}%) 内容匹配...",
                end="",
                file=sys.stderr,
            )

            # 尝试从索引获取内容
            cached = file_index.get_file_cache(filepath)

            if cached:
                # 使用缓存内容
                content = cached.get("content", "")
            else:
                # 解析文件并更新索引
                result = read_file_content(filepath)
                if result["error"]:
                    continue
                content = result["content"]
                file_index.update_file(filepath, content)

            # 正则匹配
            if content:
                matches = regex.findall(content)
                if matches:
                    content_results.append({
                        "filepath": filepath,
                        "matches": matches[:10],
                        "match_count": len(matches),
                        "match_type": "content"
                    })

        # 保存索引（含所有已索引文件）
        print("\n[索引] 保存索引...", file=sys.stderr)
        with file_lock(lock_file, timeout=5.0):
            file_index.save()
    else:
        # 不使用索引，直接解析
        for index, filepath in enumerate(files):
            progress = int((index + 1) / total_files * 100) if total_files > 0 else 0
            print(
                f"\r[阶段2] {index + 1}/{total_files} ({progress}%) 内容匹配...",
                end="",
                file=sys.stderr,
            )

            result = search_in_file(filepath, pattern)
            if result is not None:
                result["match_type"] = "content"
                content_results.append(result)

    print("", file=sys.stderr)  # 换行

    # 合并结果：文件名匹配优先
    # 去重：如果一个文件既匹配文件名又匹配内容，只保留文件名匹配结果
    filename_paths = {r["filepath"] for r in filename_results}
    filtered_content_results = [
        r for r in content_results if r["filepath"] not in filename_paths
    ]

    all_results = filename_results + filtered_content_results

    return {
        "results": all_results,
        "total_files": total_files,
        "matched_files": len(all_results),
        "filename_matches": len(filename_results),
        "content_matches": len(filtered_content_results),
        "errors": errors,
    }


def build_arg_parser():
    """
    构建命令行参数解析器。

    @returns {argparse.ArgumentParser} 参数解析器实例
    """
    parser = argparse.ArgumentParser(
        description="文件内容探索引擎 - 在 PDF/DOCX/Excel 文件中搜索关键词或读取文件完整内容"
    )

    subparsers = parser.add_subparsers(dest="command", help="可用的子命令")

    # search 子命令 - 关键词搜索
    search_parser = subparsers.add_parser("search", help="按关键词搜索文件内容，返回匹配的文件路径")
    search_parser.add_argument(
        "--dirs",
        required=True,
        help="要搜索的目录列表，用逗号分隔",
    )
    search_parser.add_argument(
        "--pattern",
        required=True,
        help="搜索用的正则表达式",
    )
    search_parser.add_argument(
        "--extensions",
        default=".pdf,.docx,.xlsx,.xls",
        help="要搜索的文件扩展名，用逗号分隔（默认: .pdf,.docx,.xlsx,.xls）",
    )
    search_parser.add_argument(
        "--index",
        action="store_true",
        default=True,
        help="使用索引加速搜索（默认启用）",
    )

    # read 子命令 - 按路径读取文件内容
    read_parser = subparsers.add_parser(
        "read",
        help="根据文件路径读取文件完整内容，支持单个或多个文件",
    )
    read_group = read_parser.add_mutually_exclusive_group(required=True)
    read_group.add_argument(
        "--files",
        help="文件路径列表，用逗号分隔（例如: /a.pdf,/b.docx）",
    )
    read_group.add_argument(
        "--files-json",
        dest="files_json",
        help='文件路径的 JSON 数组（例如: \'["/a.pdf", "/b.docx"]\'）',
    )

    # check 子命令 - 检查依赖
    subparsers.add_parser("check", help="检查 Python 依赖是否已安装")

    # index 子命令 - 构建索引
    subparsers.add_parser("index", help="构建文件内容索引")

    # index-status 子命令 - 查看索引状态
    subparsers.add_parser("index-status", help="查看索引状态")

    # index-rebuild 子命令 - 强制重建索引
    subparsers.add_parser("index-rebuild", help="强制重建索引")

    return parser


def handle_search_command(args):
    """
    处理 search 子命令的逻辑。

    @param {argparse.Namespace} args - 解析后的命令行参数
    """
    directories = [d.strip() for d in args.dirs.split(",") if d.strip()]
    pattern = args.pattern
    extensions = [e.strip() for e in args.extensions.split(",") if e.strip()]

    # 验证正则表达式是否合法
    try:
        re.compile(pattern)
    except re.error as e:
        error_result = {"error": f"无效的正则表达式: {str(e)}"}
        print(json.dumps(error_result, ensure_ascii=False))
        sys.exit(1)

    # 索引文件路径（统一存放在 ~/.fileIndex/ 目录）
    index_path = None
    if hasattr(args, 'index') and args.index:
        index_root = os.path.expanduser("~/.fileIndex")
        os.makedirs(index_root, exist_ok=True)
        index_path = os.path.join(index_root, "file_index.db")

    result = search_in_directories(directories, pattern, extensions, index_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def handle_read_command(args):
    """
    处理 read 子命令的逻辑：根据文件路径列表读取并返回完整文件内容。

    @param {argparse.Namespace} args - 解析后的命令行参数
    """
    filepaths = []

    if args.files_json:
        # 解析 JSON 数组格式
        try:
            raw = json.loads(args.files_json)
            if not isinstance(raw, list):
                print(json.dumps({"error": "--files-json 必须是一个 JSON 数组"}, ensure_ascii=False))
                sys.exit(1)
            filepaths = [str(p).strip() for p in raw if str(p).strip()]
        except json.JSONDecodeError as e:
            print(json.dumps({"error": f"JSON 解析失败: {str(e)}"}, ensure_ascii=False))
            sys.exit(1)
    else:
        # 逗号分隔格式
        filepaths = [p.strip() for p in args.files.split(",") if p.strip()]

    if not filepaths:
        print(json.dumps({"error": "未提供任何文件路径"}, ensure_ascii=False))
        sys.exit(1)

    result = read_multiple_files(filepaths)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def handle_check_command():
    """
    处理 check 子命令的逻辑，检查并输出依赖安装状态。
    """
    deps = check_dependencies()
    print(json.dumps(deps, ensure_ascii=False))


def handle_index_command(args):
    """
    处理 index 子命令：构建文件内容索引（增量更新）
    """
    # 从配置文件读取目录和扩展名（统一路径：~/.fileIndex/）
    index_root = os.path.expanduser("~/.fileIndex")
    os.makedirs(index_root, exist_ok=True)

    config_path = os.path.join(index_root, "file_explorer_config.json")
    index_path = os.path.join(index_root, "file_index.db")

    # 配置文件不存在，使用默认配置
    if not os.path.exists(config_path):
        print(f"[索引] 配置文件不存在，使用默认配置（全部目录 + 全部文件类型）", file=sys.stderr)

        # 扫描 HOME 下所有一级子目录
        home = os.path.expanduser("~")
        directories = []
        for item in os.listdir(home):
            item_path = os.path.join(home, item)
            # 只添加目录，排除隐藏目录和 ~/.fileIndex/
            if os.path.isdir(item_path) and not item.startswith('.') and item_path != index_root:
                directories.append(item_path)

        # 默认扩展名
        extensions = [".pdf", ".docx", ".xlsx", ".xls"]
    else:
        # 读取配置
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)

        directories = config.get("directories", [])
        extensions = config.get("extensions", [])

    if not directories:
        print(json.dumps({"error": "没有有效的搜索目录"}, ensure_ascii=False))
        sys.exit(1)

    print(f"[索引] 开始构建索引...", file=sys.stderr)
    print(f"[索引] 目录: {', '.join(directories)}", file=sys.stderr)
    print(f"[索引] 文件类型: {', '.join(extensions)}", file=sys.stderr)

    # 收集所有文件
    files = collect_files(directories, extensions)
    total_files = len(files)
    print(f"[索引] 共发现 {total_files} 个文件", file=sys.stderr)

    # 初始化锁文件路径
    lock_file = os.path.join(index_root, ".lock")

    # 初始化索引
    file_index = FileIndex(index_path, lock_file)

    # 统计：新增、更新、跳过
    stats = {
        "new": 0,
        "updated": 0,
        "skipped": 0,
        "errors": 0
    }

    # 增量更新：只处理变化的文件
    for index, filepath in enumerate(files):
        progress = int((index + 1) / total_files * 100) if total_files > 0 else 0
        print(
            f"\r[索引] {index + 1}/{total_files} ({progress}%) {os.path.basename(filepath)}",
            end="",
            file=sys.stderr,
        )

        # 检查缓存有效性
        cached = file_index.get_file_cache(filepath)
        if cached:
            # 缓存有效，跳过
            stats["skipped"] += 1
            continue

        # 缓存无效，需要解析
        result = read_file_content(filepath)
        if result["error"] is None:
            # 判断是新增还是更新（SQLite: 通过查询判断是否存在）
            existing = file_index.get_file_cache(filepath)
            if existing:
                stats["updated"] += 1
            else:
                stats["new"] += 1
            file_index.update_file(filepath, result["content"])
        else:
            stats["errors"] += 1

    print("\n[索引] 保存索引...", file=sys.stderr)
    with file_lock(lock_file, timeout=5.0):
        file_index.save()

    # 输出统计
    print(f"\n[统计] 新增: {stats['new']} | 更新: {stats['updated']} | 跳过: {stats['skipped']} | 错误: {stats['errors']}", file=sys.stderr)

    # 输出结果
    status = file_index.get_status()
    print(json.dumps({
        "success": True,
        "total_files": status["total_files"],
        "total_size_mb": status["total_size_mb"],
        "index_file": status["index_file"],
        "stats": stats
    }, ensure_ascii=False, indent=2))


def handle_index_status_command():
    """
    处理 index-status 子命令：查看索引状态
    """
    index_root = os.path.expanduser("~/.fileIndex")
    index_path = os.path.join(index_root, "file_index.db")

    file_index = FileIndex(index_path)
    status = file_index.get_status()

    print(json.dumps(status, ensure_ascii=False, indent=2))


def handle_index_rebuild_command():
    """
    处理 index-rebuild 子命令：强制重建索引
    """
    index_root = os.path.expanduser("~/.fileIndex")
    index_path = os.path.join(index_root, "file_index.db")

    # 清空现有索引
    file_index = FileIndex(index_path)
    file_index.clear()

    print("[索引] 已清空现有索引，开始重建...", file=sys.stderr)

    # 调用 index 命令重新构建
    handle_index_command(None)


def main():
    """
    主入口函数，解析命令行参数并分发到对应的处理函数。
    """
    parser = build_arg_parser()
    args = parser.parse_args()

    if args.command == "search":
        handle_search_command(args)
    elif args.command == "read":
        handle_read_command(args)
    elif args.command == "check":
        handle_check_command()
    elif args.command == "index":
        handle_index_command(args)
    elif args.command == "index-status":
        handle_index_status_command()
    elif args.command == "index-rebuild":
        handle_index_rebuild_command()
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
